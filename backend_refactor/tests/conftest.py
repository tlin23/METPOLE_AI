import json
import pytest
from backend_refactor.models.content_chunk import ContentChunk
from unittest.mock import patch, Mock
from docx import Document


# ===== Basic Test Infrastructure =====


@pytest.fixture
def temp_dir(tmp_path):
    """Create a temporary directory for test output.

    Returns:
        Path: A temporary directory path for test files.
    """
    return tmp_path


# ===== Content Chunk Fixtures =====


@pytest.fixture
def sample_chunks():
    """Create sample content chunks for testing.

    Returns:
        List[ContentChunk]: A list of content chunks with various properties.
        Includes chunks with different file types, page numbers, and content.
    """
    return [
        ContentChunk(
            chunk_id="chunk_1",
            file_name="test1",
            file_ext=".html",
            page_number=1,
            text_content="Test content 1",
            document_title="Test Document 1",
        ),
        ContentChunk(
            chunk_id="chunk_2",
            file_name="test2",
            file_ext=".pdf",
            page_number=2,
            text_content="Test content 2",
            document_title="Test Document 2",
        ),
    ]


@pytest.fixture
def valid_content_chunk():
    """Create a valid content chunk for testing.

    Returns:
        ContentChunk: A valid content chunk with all required fields.
    """
    return ContentChunk(
        chunk_id="test123",
        file_name="test_file",
        file_ext=".txt",
        page_number=1,
        text_content="This is a test content",
        document_title="Test Document",
    )


@pytest.fixture
def invalid_content_chunk_data():
    """Create invalid content chunk data for testing.

    Returns:
        Dict: A dictionary with invalid data for ContentChunk validation.
        Contains an invalid type for page_number to test validation.
    """
    return {
        "chunk_id": "test123",
        "file_name": "test_file",
        "file_ext": "txt",
        "page_number": "not_an_integer",  # Invalid type
        "text_content": "This is a test content",
    }


# ===== File Fixtures =====


@pytest.fixture
def test_html_file(temp_dir):
    """Create a test HTML file with basic content.

    Returns:
        Path: Path to a test HTML file with basic structure.
    """
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Test Page</title>
    </head>
    <body>
        <h1>Main Heading</h1>
        <p>This is a test paragraph.</p>
    </body>
    </html>
    """
    file_path = temp_dir / "test.html"
    file_path.write_text(html_content)
    return file_path


@pytest.fixture
def complex_html_file(temp_dir):
    """Create a test HTML file with complex content.

    Returns:
        Path: Path to a test HTML file with complex structure and content.
    """
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Complex Test Page</title>
    </head>
    <body>
        <header>Header Content</header>
        <nav>Navigation</nav>
        <main>
            <article>
                <h1>Article Title</h1>
                <p>First paragraph of the article.</p>
                <h2>Subsection</h2>
                <p>Second paragraph with more content.</p>
                <p>Third paragraph with additional information.</p>
            </article>
            <section>
                <h3>Another Section</h3>
                <p>Content in another section.</p>
            </section>
        </main>
        <footer>Footer Content</footer>
    </body>
    </html>
    """
    file_path = temp_dir / "complex.html"
    file_path.write_text(html_content)
    return file_path


@pytest.fixture
def duplicate_html_file(temp_dir):
    """Create a test HTML file with duplicate content.

    Returns:
        Path: Path to a test HTML file with duplicate paragraphs.
    """
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Duplicate Test</title>
    </head>
    <body>
        <p>This is a duplicate paragraph.</p>
        <p>This is a duplicate paragraph.</p>
        <p>This is a different paragraph.</p>
        <p>This is a different paragraph.</p>
        <p>This is a unique paragraph.</p>
    </body>
    </html>
    """
    file_path = temp_dir / "duplicate.html"
    file_path.write_text(html_content)
    return file_path


@pytest.fixture
def test_docx_file(temp_dir):
    """Create a test DOCX file with basic content.

    Returns:
        Path: Path to a test DOCX file with basic structure.
    """
    doc = Document()
    doc.core_properties.title = "Test Document"

    # Add heading
    doc.add_heading("Main Heading", level=1)

    # Add paragraph
    doc.add_paragraph("This is a test paragraph.")

    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 1"
    table.cell(0, 1).text = "Cell 2"
    table.cell(1, 0).text = "Cell 3"
    table.cell(1, 1).text = "Cell 4"

    file_path = temp_dir / "test.docx"
    doc.save(file_path)
    return file_path


@pytest.fixture
def complex_docx_file(temp_dir):
    """Create a test DOCX file with complex content.

    Returns:
        Path: Path to a test DOCX file with complex structure and content.
    """
    doc = Document()
    doc.core_properties.title = "Complex Test Document"

    # Add main heading
    doc.add_heading("Document Title", level=1)

    # Add paragraphs
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.add_paragraph("Third paragraph with additional information.")

    # Add subsection
    doc.add_heading("Another Section", level=2)
    doc.add_paragraph("Content in another section.")

    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Table Cell 1"
    table.cell(0, 1).text = "Table Cell 2"
    table.cell(1, 0).text = "Table Cell 3"
    table.cell(1, 1).text = "Table Cell 4"

    file_path = temp_dir / "complex.docx"
    doc.save(file_path)
    return file_path


@pytest.fixture
def duplicate_docx_file(temp_dir):
    """Create a test DOCX file with duplicate content.

    Returns:
        Path: Path to a test DOCX file with duplicate paragraphs.
    """
    doc = Document()
    doc.core_properties.title = "Duplicate Test Document"

    # Add duplicate paragraphs
    doc.add_paragraph("This is a duplicate paragraph.")
    doc.add_paragraph("This is a duplicate paragraph.")  # Duplicate
    doc.add_paragraph("This is a different paragraph.")
    doc.add_paragraph("This is a duplicate paragraph.")  # Duplicate
    doc.add_paragraph("This is a unique paragraph.")

    file_path = temp_dir / "duplicate.docx"
    doc.save(file_path)
    return file_path


@pytest.fixture
def test_pdf_file(temp_dir):
    """Create a test PDF file with basic content.

    Returns:
        Path: Path to a test PDF file with basic structure.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    file_path = temp_dir / "test.pdf"
    c = canvas.Canvas(str(file_path), pagesize=letter)

    # Set title
    c.setTitle("Test PDF")

    # Add content
    c.drawString(100, 750, "First paragraph of the test document.")
    c.drawString(100, 730, "Second paragraph with more content.")

    c.save()
    return file_path


@pytest.fixture
def complex_pdf_file(temp_dir):
    """Create a test PDF file with complex content.

    Returns:
        Path: Path to a test PDF file with complex structure and content.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    file_path = temp_dir / "complex.pdf"
    c = canvas.Canvas(str(file_path), pagesize=letter)

    # Set title
    c.setTitle("Complex Test PDF")

    # First page content
    c.drawString(100, 750, "First paragraph of the document.")
    c.drawString(100, 730, "Second paragraph with more content.")
    c.drawString(100, 710, "Third paragraph with additional information.")

    # Add second page
    c.showPage()
    c.drawString(100, 750, "Content on the second page.")
    c.drawString(100, 730, "Final paragraph of the document.")

    c.save()
    return file_path


@pytest.fixture
def duplicate_pdf_file(temp_dir):
    """Create a test PDF file with duplicate content.

    Returns:
        Path: Path to a test PDF file with duplicate paragraphs.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    file_path = temp_dir / "duplicate.pdf"
    c = canvas.Canvas(str(file_path), pagesize=letter)

    # Set title
    c.setTitle("Duplicate Test PDF")

    # Add content with duplicates
    c.drawString(100, 750, "This is a duplicate paragraph.")
    c.drawString(100, 730, "This is a duplicate paragraph.")  # Duplicate
    c.drawString(100, 710, "This is a different paragraph.")
    c.drawString(100, 690, "This is a duplicate paragraph.")  # Duplicate
    c.drawString(100, 670, "This is a unique paragraph.")

    c.save()
    return file_path


# ===== JSON File Fixtures =====


@pytest.fixture
def sample_json_file(sample_chunks, temp_dir):
    """Create a temporary JSON file with sample chunks.

    Returns:
        Path: Path to a JSON file containing serialized content chunks.
    """
    json_path = temp_dir / "test_chunks.json"
    chunks_data = [chunk.model_dump() for chunk in sample_chunks]
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(chunks_data, f)
    return json_path


@pytest.fixture
def multiple_json_files(sample_chunks, temp_dir):
    """Create multiple JSON files with different chunks.

    Returns:
        List[Path]: List of paths to JSON files, each containing different chunks.
    """
    json_path1 = temp_dir / "chunks1.json"
    json_path2 = temp_dir / "chunks2.json"

    with open(json_path1, "w", encoding="utf-8") as f:
        json.dump([sample_chunks[0].model_dump()], f)
    with open(json_path2, "w", encoding="utf-8") as f:
        json.dump([sample_chunks[1].model_dump()], f)

    return [json_path1, json_path2]


# ===== Mock Fixtures =====


@pytest.fixture
def mock_parser_class():
    """Create a mock parser class for testing.

    Returns:
        Type: A mock parser class that implements the parser interface.
    """

    class MockParserClass:
        def __init__(self):
            pass

        def parse(self, file_path):
            return [
                ContentChunk(
                    chunk_id="chunk_1",
                    file_name="test",
                    file_ext=".html",
                    page_number=1,
                    text_content="Test content",
                )
            ]

    return MockParserClass


@pytest.fixture
def parser_map_patch(mock_parser_class):
    """Create a patch for PARSER_MAP with mock parser.

    Returns:
        Mock: A patch object for PARSER_MAP with mock parser.
    """
    return patch(
        "backend_refactor.pipeline.pipeline_orchestration.PARSER_MAP",
        {".html": mock_parser_class},
    )


@pytest.fixture
def mock_collection():
    """Create a mock ChromaDB collection.

    Returns:
        Mock: A mock ChromaDB collection with add method.
    """
    collection = Mock()
    collection.add = Mock()
    return collection


@pytest.fixture
def mock_client(mock_collection):
    """Create a mock ChromaDB client.

    Returns:
        Mock: A mock ChromaDB client with get_or_create_collection method.
    """
    client = Mock()
    client.get_or_create_collection = Mock(return_value=mock_collection)
    return client


@pytest.fixture
def mock_response():
    """Create a mock response object for web crawler tests.

    Returns:
        Mock: A mock response object with HTML content and raise_for_status method.
    """
    mock = Mock()
    mock.text = """
    <html>
        <body>
            <a href="https://example.com/page1">Link 1</a>
            <a href="https://example.com/page2">Link 2</a>
            <a href="https://other.com/page3">Link 3</a>
        </body>
    </html>
    """
    mock.raise_for_status = Mock()
    return mock


# ===== Directory Structure Fixtures =====


@pytest.fixture
def sample_files(temp_dir):
    """Create sample files for local crawler tests.

    Returns:
        Path: Path to a directory containing various test files.
    """
    input_dir = temp_dir / "input"
    input_dir.mkdir()

    # Create files with different extensions
    (input_dir / "test1.txt").write_text("Content 1")
    (input_dir / "test2.txt").write_text("Content 2")
    (input_dir / "test3.pdf").write_text("Content 3")
    (input_dir / "test4.doc").write_text("Content 4")

    return input_dir


@pytest.fixture
def test_directory_structure(temp_dir):
    """Create a test directory structure with various file types.

    Returns:
        Path: Path to a root directory containing a complex file structure.
    """
    root = temp_dir / "test_root"
    root.mkdir()

    # Create subdirectories
    docs = root / "docs"
    images = root / "images"
    docs.mkdir()
    images.mkdir()

    # Create files in root
    (root / "index.html").write_text("<html><body>Index</body></html>")
    (root / "config.txt").write_text("config content")

    # Create files in docs
    (docs / "manual.pdf").write_bytes(b"%PDF-1.4\n%Test PDF content")
    (docs / "guide.docx").write_bytes(b"PK\x03\x04\x14\x00\x00\x00\x08\x00")

    # Create files in images
    (images / "logo.png").write_bytes(b"PNG\r\n\x1a\n")
    (images / "icon.ico").write_bytes(b"ICO")

    return root


# ===== URL and Network Fixtures =====


@pytest.fixture
def test_urls():
    """Create test URLs for testing.

    Returns:
        Dict: Dictionary containing lists of valid and invalid URLs.
    """
    return {
        "valid": [
            "https://example.com",
            "http://localhost:8000",
            "http://example.com/path?query=value",
        ],
        "invalid": [
            "not-a-url",
            "ftp://example.com",
            "file:///path/to/file",
            "",
            "http://",
        ],
    }


# ===== Error Message Fixtures =====


@pytest.fixture
def test_error_messages():
    """Create common test error messages.

    Returns:
        Dict: Dictionary containing common error message templates.
    """
    return {
        "file_not_found": "File not found: {path}",
        "invalid_file": "Invalid file format: {file}",
        "network_error": "Network error: {error}",
        "parsing_error": "Error parsing file: {file}",
        "embedding_error": "Error embedding chunks: {error}",
    }
