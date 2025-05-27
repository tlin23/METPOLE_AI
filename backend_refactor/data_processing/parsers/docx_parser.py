from pathlib import Path
import re
import hashlib
from typing import List, Set
from docx import Document
from ftfy import fix_text
from .base import BaseParser
from ..models.content_chunk import ContentChunk
from ...logger.logging_config import get_logger

logger = get_logger("parsers.docx")


def clean_text(text: str) -> str:
    """Clean and normalize text by fixing encoding, stripping special characters, and collapsing whitespace."""
    text = fix_text(text)
    # Replace zero-width spaces with regular spaces
    text = text.replace("\u200b", " ")
    # Remove other invisible characters
    text = re.sub(r"[\u200e\u200f\u202a-\u202e]", "", text)
    text = (
        text.replace(
            """, '"')
        .replace(""",
            '"',
        )
        .replace("'", "'")
        .replace("'", "'")
        .replace("–", "-")
        .replace("—", "-")
        .replace("…", "...")
    )
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize(text: str) -> str:
    """Normalize text for hashing by converting to lowercase and removing non-word characters."""
    text = text.lower()
    text = re.sub(r"\W+", " ", text)
    return " ".join(text.split())


def hash_id(text: str) -> str:
    """Generate a unique hash ID for a text chunk."""
    normalized = normalize(text)
    return "chunk_" + hashlib.md5(normalized.encode("utf-8")).hexdigest()


class DOCXParser(BaseParser):
    def parse(self, file_path: Path) -> List[ContentChunk]:
        """
        Parse a DOCX file and convert its contents into a list of ContentChunk objects.

        The parser:
        - Extracts meaningful content while preserving document structure
        - Handles paragraphs, headings, and tables
        - Extracts metadata like document title and properties
        - Cleans and normalizes text
        - Generates unique chunk IDs
        - Prevents duplicate chunks

        Args:
            file_path: Path to the DOCX file to parse

        Returns:
            List of ContentChunk objects representing the parsed content

        Raises:
            ValueError: If the file is not a valid DOCX file
            IOError: If the file cannot be read
        """
        logger.info(f"Starting to parse DOCX file: {file_path}")

        if not file_path.exists():
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            raise IOError(error_msg)

        if file_path.suffix.lower() != ".docx":
            error_msg = f"Not a DOCX file: {file_path}"
            logger.error(error_msg)
            raise ValueError(error_msg)

        try:
            doc = Document(file_path)
            chunks = []
            seen_chunk_ids: Set[str] = set()
            current_heading = None
            paragraph_count = 0
            table_count = 0

            # Extract document title from properties
            document_title = doc.core_properties.title or file_path.stem
            logger.debug(f"Extracted document title: {document_title}")

            # Process paragraphs and headings
            logger.debug("Processing paragraphs and headings")
            for para in doc.paragraphs:
                text = clean_text(para.text)
                if not text:
                    continue

                # Check if paragraph is a heading
                if para.style.name.startswith("Heading"):
                    current_heading = text
                    logger.debug(f"Found heading: {current_heading}")
                    continue

                # Skip very short chunks
                if len(text) < 20:
                    logger.debug(f"Skipping short paragraph: {text[:50]}...")
                    continue

                # Combine heading with content if available
                if current_heading:
                    text = f"{current_heading}\n{text}"
                    current_heading = None  # Reset heading after using it

                # Generate chunk ID and check for duplicates
                chunk_id = hash_id(file_path.stem + text)
                if chunk_id in seen_chunk_ids:
                    logger.debug(f"Skipping duplicate paragraph chunk: {chunk_id}")
                    continue
                seen_chunk_ids.add(chunk_id)

                # Create chunk
                chunk = ContentChunk(
                    chunk_id=chunk_id,
                    file_name=file_path.stem,
                    file_ext=file_path.suffix[1:],
                    page_number=1,  # DOCX doesn't have explicit page numbers
                    text_content=text,
                    document_title=document_title,
                )
                chunks.append(chunk)
                paragraph_count += 1

            # Process tables
            logger.debug("Processing tables")
            for table_idx, table in enumerate(doc.tables, 1):
                table_text = []
                for row in table.rows:
                    row_text = [clean_text(cell.text) for cell in row.cells]
                    table_text.append(" | ".join(row_text))

                table_content = "\n".join(table_text)
                if not table_content.strip():
                    logger.debug(f"Skipping empty table {table_idx}")
                    continue

                # Generate chunk ID and check for duplicates
                chunk_id = hash_id(file_path.stem + table_content)
                if chunk_id in seen_chunk_ids:
                    logger.debug(f"Skipping duplicate table chunk: {chunk_id}")
                    continue
                seen_chunk_ids.add(chunk_id)

                # Create chunk for table
                chunk = ContentChunk(
                    chunk_id=chunk_id,
                    file_name=file_path.stem,
                    file_ext=file_path.suffix[1:],
                    page_number=1,
                    text_content=f"Table {table_idx}:\n{table_content}",
                    document_title=document_title,
                )
                chunks.append(chunk)
                table_count += 1

            logger.info(
                f"Successfully parsed {paragraph_count} paragraphs and {table_count} tables from {file_path}"
            )
            return chunks

        except Exception as e:
            error_msg = f"Failed to parse DOCX file: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg)
