import time
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from backend.configer.logging_config import get_logger

logger = get_logger("pipeline.processors.docx")


def clean_text(text: str) -> str:
    """Clean and normalize text content."""
    return " ".join(text.split()).strip()


class DOCXProcessor:
    """Processor for DOCX documents."""

    def process(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process a DOCX document and return structured content blocks."""
        logger.info(f"Starting to process DOCX file: {file_path.name}")
        start_time = time.time()

        try:
            # Load document
            logger.debug(f"Loading document: {file_path}")
            doc = Document(file_path)
            logger.info(
                f"Document loaded successfully. Found {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
            )

            # Extract content
            structured_blocks = []
            current_section = "Document Content"
            logger.debug(f"Initial section set to: {current_section}")

            # Add document properties
            logger.debug("Extracting document properties")
            properties = self._extract_document_properties(doc)
            if properties:
                logger.info("Document properties extracted successfully")
                structured_blocks.append(
                    {
                        "type": "properties",
                        "section": "Document Properties",
                        "content": properties,
                    }
                )

            # Process main content
            logger.debug("Starting to process main content")
            element_count = {"paragraphs": 0, "tables": 0, "lists": 0, "headings": 0}

            for element in doc.element.body:
                if isinstance(element, CT_P):  # Paragraph
                    element_count["paragraphs"] += 1
                    paragraph = Paragraph(element, doc)

                    # Handle headings
                    if self._is_heading(paragraph):
                        element_count["headings"] += 1
                        current_section = clean_text(paragraph.text)
                        logger.debug(f"Found heading: {current_section}")
                        continue

                    # Handle lists
                    if paragraph.style.name.startswith("List"):
                        element_count["lists"] += 1
                        list_text = self._extract_list_item(paragraph)
                        if list_text:
                            logger.debug(f"Found list item: {list_text[:50]}...")
                            structured_blocks.append(
                                {
                                    "type": "list",
                                    "section": current_section,
                                    "content": list_text,
                                }
                            )
                        continue

                    # Handle regular text
                    text = clean_text(paragraph.text)
                    if text:  # Skip empty paragraphs
                        logger.debug(
                            f"Found text in section '{current_section}': {text[:50]}..."
                        )
                        structured_blocks.append(
                            {
                                "type": "text",
                                "section": current_section,
                                "content": text,
                            }
                        )

                elif isinstance(element, CT_Tbl):  # Table
                    element_count["tables"] += 1
                    table = Table(element, doc)
                    table_text = self._extract_table_text(table)
                    if table_text:
                        logger.debug(
                            f"Found table in section '{current_section}' with {len(table.rows)} rows"
                        )
                        structured_blocks.append(
                            {
                                "type": "table",
                                "section": current_section,
                                "content": table_text,
                            }
                        )

            # Log processing summary
            logger.info("Document processing completed successfully")
            logger.info("Processing summary:")
            logger.info(f"- Total paragraphs processed: {element_count['paragraphs']}")
            logger.info(f"- Total tables processed: {element_count['tables']}")
            logger.info(f"- Total lists found: {element_count['lists']}")
            logger.info(f"- Total headings found: {element_count['headings']}")
            logger.info(f"- Total content blocks created: {len(structured_blocks)}")
            logger.info(f"Processing took {time.time() - start_time:.2f} seconds")

            return structured_blocks

        except Exception as e:
            logger.error(
                f"Error processing document {file_path}: {str(e)}", exc_info=True
            )
            raise

    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Check if a paragraph is a heading."""
        is_heading = paragraph.style.name.startswith("Heading")
        if is_heading:
            logger.debug(f"Identified heading: {paragraph.style.name}")
        return is_heading

    def _extract_list_item(self, paragraph: Paragraph) -> Optional[str]:
        """Extract text from a list item with basic formatting."""
        text = clean_text(paragraph.text)
        if not text:
            logger.debug("Empty list item found, skipping")
            return None

        # Get list level
        p = paragraph._p
        num_pr = p.pPr.numPr if p.pPr else None
        if not num_pr:
            logger.debug("List item without numbering properties found, skipping")
            return None

        level = num_pr.ilvl.val if num_pr.ilvl else 0
        logger.debug(f"Processing list item at level {level}")

        # Format based on list style
        if paragraph.style.name.startswith("List Bullet"):
            prefix = "• " * (level + 1)
            logger.debug("Formatted as bullet list")
        else:  # Numbered list
            prefix = f"{level + 1}. "
            logger.debug("Formatted as numbered list")

        return f"{prefix}{text}"

    def _extract_table_text(self, table: Table) -> str:
        """Extract text from table in a structured format."""
        logger.debug(
            f"Processing table with {len(table.rows)} rows and {len(table.columns)} columns"
        )
        table_text = []

        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
                logger.debug(f"Processed row {row_idx + 1}: {row_text[0][:50]}...")

        result = "\n".join(table_text) if table_text else ""
        logger.debug(f"Table processing complete. Extracted {len(table_text)} rows")
        return result

    def _extract_document_properties(self, doc: Document) -> Dict[str, Any]:
        """Extract basic document properties."""
        logger.debug("Starting document properties extraction")
        core_properties = doc.core_properties
        properties = {}

        # Basic properties
        if core_properties.title:
            properties["title"] = core_properties.title
            logger.debug(f"Found title: {core_properties.title}")
        if core_properties.author:
            properties["author"] = core_properties.author
            logger.debug(f"Found author: {core_properties.author}")
        if core_properties.created:
            properties["created"] = core_properties.created.isoformat()
            logger.debug(f"Found creation date: {core_properties.created}")
        if core_properties.modified:
            properties["modified"] = core_properties.modified.isoformat()
            logger.debug(f"Found modification date: {core_properties.modified}")

        # Document statistics
        properties["statistics"] = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        logger.debug(f"Document statistics: {properties['statistics']}")

        return properties
